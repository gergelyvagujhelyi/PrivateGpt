from __future__ import annotations

import io

import pytest
from docx import Document

from src.extract import UnsupportedFormat, extract_text


def test_extracts_utf8_text():
    assert extract_text("readme.txt", "text/plain", b"hello") == "hello"


def test_extracts_markdown_via_suffix():
    text = extract_text("page.md", "application/octet-stream", b"# title\n\nbody")
    assert "title" in text and "body" in text


def test_extracts_docx_from_bytes():
    buf = io.BytesIO()
    d = Document()
    d.add_paragraph("first paragraph")
    d.add_paragraph("second paragraph")
    d.save(buf)
    text = extract_text("doc.docx", "", buf.getvalue())
    assert "first paragraph" in text and "second paragraph" in text


def test_unsupported_format_raises():
    with pytest.raises(UnsupportedFormat):
        extract_text("image.png", "image/png", b"")
